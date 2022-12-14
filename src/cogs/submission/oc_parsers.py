# Copyright 2022 Vioshim
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#      https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.


from abc import ABCMeta, abstractmethod
from asyncio import gather
from contextlib import suppress
from enum import Enum
from io import BytesIO
from itertools import chain
from typing import Any, Optional

from discord import File, Message
from docx import Document
from docx.document import Document as DocumentType
from jishaku.codeblocks import codeblock_converter
from yaml import safe_load
from yaml.scanner import ScannerError

from src.structures.bot import CustomBot
from src.utils.doc_reader import DriveFormat, docs_aioreader
from src.utils.functions import yaml_handler
from src.utils.matches import DATA_FINDER, G_DOCUMENT, REGEX_URL

__all__ = ("OCParser", "ParserMethods")


PLACEHOLDER_NAMES = {
    "Name": "name",
    "Names": "name",
    "Age": "age",
    "Gender": "gender",
    "Ability": "ability",
    "Abilities": "abilities",
    "Pronoun": "pronoun",
    "Moveset": "moveset",
    "Backstory": "backstory",
    "Personality": "personality",
    "Types": "types",
    "Additional Information": "extra",
    # Backwards compatibilty
    "Chimera": "chimera",
    "F. Base": "base",
    "F. Species": "fakemon",
    "Mega Base": "mega",
    "Mega Species": "mega",
    "Paradox Base": "paradox",
    "Variant": "variant",
    "Artist": "artist",
    "Website": "website",
    # Species
    "Species": "species",
    "Chimera Species": "chimera",
    "Base Species": "base",
    "Fakemon Base": "base",
    "Variant Species": "variant",
    "Fakemon Species": "fakemon",
    "Paradox Species": "paradox",
}
PLACEHOLDER_DEFAULTS = {
    "name": "OC's Name",
    "age": "OC's Age",
    "species": "OC's Species",
    "gender": "OC's Gender",
    "ability": "OC's Ability",
    "abilities": "OC's Abilities",
    "moveset": "OC's Moveset",
    "types": "OC's Types",
    "pronoun": "OC's Preferred Pronoun",
    "backstory": "Character's backstory",
    "personality": "Character's personality",
    "extra": "Character's extra information",
    "fakemon": "OC's Fakemon Species",
    "base": "OC's Base Species",
    "mega": "OC's Base Species",
    "chimera": "OC's Chimera Species",
    "variant": "OC's Variant Species",
    "paradox": "OC's Paradox Species",
    "artist": "Artist's Name",
    "website": "Art's Website",
}
PLACEHOLDER_SP = {
    "What is it Called?": "name",
    "How is it Called?": "name",
    "How did they obtain it?": "origin",
    "What does the Special Ability do?": "description",
    "What does the Unique Trait do?": "description",
    "How does it make the character's life easier?": "pros",
    "How does it make the character's life harder?": "cons",
}
PLACEHOLDER_SP_DEFAULTS = {
    "name": "Here you can describe how the unique trait is called, either how it's referenced OOC or in RP.",
    "origin": "Here you can describe how the trait was obtained.",
    "description": "Here you can describe how the trait works, what causes it to activate, the required conditions for it to happen, its cool down and such.",
    "pros": "Describe how such a trait causes pros in the character, what it means for the character to have it, and the usual things they'd use their unique trait for.",
    "cons": "Describe how such a trait causes downsides in the character, what it implies for the character to have it, and the usual things they'd try to not use their unique trait due to its limitations.",
}
PLACEHOLDER_STATS = {
    "HP": "HP",
    "Attack": "ATK",
    "Defense": "DEF",
    "Special Attack": "SPA",
    "Special Defense": "SPD",
    "Speed": "SPE",
}
IGNORE_MOVE = ["None", "Move", "Ability"]


def doc_convert(doc: DocumentType) -> dict[str, Any]:
    """Google Convereter

    Parameters
    ----------
    doc : Document
        docx Document

    Returns
    -------
    dict[str, Any]
        Info
    """
    tables = doc.tables
    tables.extend(chain(*[cell.tables for table in doc.tables for row in table.rows for cell in row.cells]))
    content_values = [str(cell.text) for table in tables for row in table.rows for cell in row.cells]

    text = [x for item in content_values if (x := item.replace("\u2019", "'").strip())]
    raw_kwargs = dict(url=getattr(doc, "url", None))

    for index, item in enumerate(text[:-1], start=1):
        next_value = text[index]
        data = f"{next_value}".title().strip() not in IGNORE_MOVE
        data &= next_value not in PLACEHOLDER_NAMES
        data &= next_value not in PLACEHOLDER_DEFAULTS.values()
        data &= next_value not in PLACEHOLDER_SP_DEFAULTS.values()
        data &= next_value not in PLACEHOLDER_STATS
        if not data:
            continue
        if argument := PLACEHOLDER_NAMES.get(item):
            if item.lower() in ["abilities", "types", "moveset"]:
                argument = item.lower()
                raw_kwargs.setdefault(argument, set())
                if isinstance(raw_kwargs[argument], str):
                    raw_kwargs[argument] = {raw_kwargs[argument]}
                values = [o for x in next_value.split(",") if (o := x.title().strip()) and o not in IGNORE_MOVE]
                raw_kwargs[argument].update(values)
            else:
                raw_kwargs[argument] = next_value
        elif element := PLACEHOLDER_SP.get(item):
            raw_kwargs.setdefault("sp_ability", {})
            raw_kwargs["sp_ability"][element] = next_value
        elif element := DATA_FINDER.match(item):
            argument = next_value.title()
            match element.groups():
                case ["Level", y]:
                    idx = int(y)
                    raw_kwargs.setdefault("movepool", {})
                    raw_kwargs["movepool"].setdefault("level", {})
                    raw_kwargs["movepool"]["level"].setdefault(idx, set())
                    info: set[str] = raw_kwargs["movepool"]["level"][idx]
                    info.update(o for x in argument.split(",") if (o := x.title().strip()) and o not in IGNORE_MOVE)
                case ["Move", _]:
                    raw_kwargs.setdefault("moveset", set())
                    if isinstance(raw_kwargs["moveset"], str):
                        raw_kwargs["moveset"] = {raw_kwargs["moveset"]}
                    raw_kwargs["moveset"].add(argument)
                case ["Ability", _]:
                    raw_kwargs.setdefault("abilities", set())
                    if isinstance(raw_kwargs["abilities"], str):
                        raw_kwargs["abilities"] = {raw_kwargs["abilities"]}
                    raw_kwargs["abilities"].add(next_value)
                case ["Species", _]:
                    raw_kwargs.setdefault("fusion", set())
                    if isinstance(raw_kwargs["fusion"], str):
                        raw_kwargs["fusion"] = {raw_kwargs["fusion"]}
                    raw_kwargs["fusion"].add(next_value)
                case ["Type", _]:
                    raw_kwargs.setdefault("types", set())
                    if isinstance(raw_kwargs["types"], str):
                        raw_kwargs["types"] = {raw_kwargs["types"]}
                    raw_kwargs["types"].add(next_value.upper())
                case [x, _]:
                    raw_kwargs.setdefault("movepool", {})
                    raw_kwargs["movepool"].setdefault(x.lower(), set())
                    if isinstance(raw_kwargs["movepool"][x.lower()], str):
                        raw_kwargs["movepool"][x.lower()] = {raw_kwargs["movepool"][x.lower()]}
                    raw_kwargs["movepool"][x.lower()].add(argument)

    try:
        if img := list(doc.inline_shapes):
            item = img[0]
            pic = item._inline.graphic.graphicData.pic
            blip = pic.blipFill.blip
            rid = blip.embed
            doc_part = doc.part
            image_part = doc_part.related_parts[rid]
            fp = BytesIO(image_part._blob)  # skipcq: PYL-W0212
            raw_kwargs["image"] = File(fp=fp, filename="image.png")
    except Exception:
        pass

    raw_kwargs.pop("artist", None)
    raw_kwargs.pop("website", None)

    return raw_kwargs


class OCParser(metaclass=ABCMeta):
    @classmethod
    @abstractmethod
    async def parse(cls, text: str | Message, bot: Optional[CustomBot] = None) -> Optional[dict[str, Any]]:
        """Parser method if possible

        Parameters
        ----------
        text : str | Message
            str or Message with information
        bot : Optional[CustomBot]
            Client instance, defaults to None

        Returns
        -------
        Optional[dict[str, Any]]
            Character information
        """


class GoogleDocsOCParser(OCParser):
    @classmethod
    async def parse(cls, text: str | Message, bot: Optional[CustomBot] = None) -> Optional[dict[str, Any]]:
        if isinstance(text, Message):
            content = text.content
        else:
            content = text
        content: str = codeblock_converter(content or "").content
        if doc_data := G_DOCUMENT.match(content):
            doc = await docs_aioreader(url := doc_data.group(1), bot.aiogoogle)
            msg_data = doc_convert(doc)
            msg_data["url"] = url
            return msg_data


class WordOCParser(OCParser):
    @classmethod
    async def parse(cls, text: str | Message, bot: Optional[CustomBot] = None) -> Optional[dict[str, Any]]:
        if isinstance(text, Message):
            for attachment in text.attachments:
                if attachment.content_type == DriveFormat.DOCX.value:
                    file = await attachment.to_file(use_cached=True)
                    doc: DocumentType = Document(file.fp)
                    if doc.tables:
                        return doc_convert(doc)
                    text = yaml_handler("\n".join(element for p in doc.paragraphs if (element := p.text.strip())))
                    with suppress(ScannerError):
                        return safe_load(text)


class DiscordOCParser(OCParser):
    @classmethod
    async def parse(cls, text: str | Message, bot: Optional[CustomBot] = None) -> Optional[dict[str, Any]]:
        if isinstance(text, Message):
            content = text.content
            images = [x for x in text.attachments if x.content_type.startswith("image/")]
        else:
            content = text
            images = None
        content = codeblock_converter(content or "").content
        if REGEX_URL.match(content) or G_DOCUMENT.match(content):
            return
        content = yaml_handler(content)
        with suppress(ScannerError):
            if isinstance(msg_data := safe_load(content), dict):
                if images:
                    msg_data["image"] = await images[0].to_file(use_cached=True)

                return msg_data


class ParserMethods(Enum):
    WORD = WordOCParser
    GOOGLEDOCS = GoogleDocsOCParser
    DISCORD = DiscordOCParser

    def __call__(self, text: str | Message, bot: Optional[CustomBot] = None):
        item: OCParser = self.value
        return item.parse(text=text, bot=bot)

    @classmethod
    async def parse(cls, text: str | Message, bot: Optional[CustomBot] = None):
        for x in await gather(*[item(text=text, bot=bot) for item in cls]):
            if x and isinstance(x, dict):
                yield x
