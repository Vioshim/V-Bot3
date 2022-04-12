# Copyright 2022 Vioshim
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#      https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from asyncio import ALL_COMPLETED, TimeoutError, to_thread, wait
from contextlib import suppress
from datetime import datetime, timedelta
from difflib import get_close_matches
from itertools import chain
from pathlib import Path
from random import choice as random_choice
from random import sample
from typing import Any, Optional, Type, Union

from aiofiles import open as aiopen
from apscheduler.enums import ConflictPolicy
from apscheduler.triggers.interval import IntervalTrigger
from asyncpg import Connection
from discord import (
    AllowedMentions,
    DiscordException,
    Embed,
    Guild,
    HTTPException,
    Interaction,
    InteractionResponse,
    Member,
    Message,
    NotFound,
    Object,
    RawMessageDeleteEvent,
    Status,
    TextChannel,
    TextStyle,
    Thread,
    User,
    Webhook,
    WebhookMessage,
    app_commands,
)
from discord.ext import commands
from discord.ui import Button, TextInput, View
from discord.utils import utcnow
from docx import Document
from docx.document import Document as DocumentType
from jishaku.codeblocks import codeblock_converter
from orjson import loads
from yaml import safe_load
from yaml.error import MarkedYAMLError

from src.pagination.boolean import BooleanView
from src.pagination.complex import Complex
from src.pagination.text_input import ModernInput
from src.structures.ability import Ability, SpAbility
from src.structures.bot import CustomBot
from src.structures.character import (
    Character,
    CharacterArg,
    doc_convert,
    fetch_all,
    oc_process,
)
from src.structures.mission import Mission
from src.structures.mon_typing import Typing
from src.structures.move import Move
from src.structures.movepool import Movepool
from src.structures.species import (
    Fakemon,
    Fusion,
    Legendary,
    Mythical,
    UltraBeast,
    Variant,
)
from src.utils.doc_reader import docs_reader
from src.utils.etc import REGISTERED_IMG, RP_CATEGORIES
from src.utils.functions import yaml_handler
from src.utils.matches import G_DOCUMENT
from src.views import (
    AbilityView,
    CharactersView,
    ImageView,
    MissionView,
    MovepoolView,
    MoveView,
    PingView,
    RPView,
    StatsView,
    SubmissionView,
)

CLAIM_MESSAGE = """
**。　　　　•　    　ﾟ　　。**
**　　.　　　.　　　.　　。　　   。　.**
** 　.　　      。             。　    .    •**
** •     RP claimable, feel free to use it.　 。　.**
**          Is assumed that everyone left**
**　 　　。　　　　ﾟ　　　.　　　　　.**
**,　　　　.　 .　　       .               。**
""".strip()


def message_validator(message: Message):
    """Function used for checking compatiblity between messages

    Parameters
    ----------
    message : Message
        Message to check
    """

    def checker(value: Message) -> bool:
        if value.webhook_id and message.channel == value.channel:
            if isinstance(content := message.content, str):
                return value.content in content
            if attachments := message.attachments:
                return len(attachments) == len(value.attachments)
        return False

    return checker


class Submission(commands.Cog):
    def __init__(self, bot: CustomBot):
        self.bot = bot
        self.ready: bool = False
        self.missions: set[Mission] = set()
        self.mission_claimers: dict[int, set[int]] = {}
        self.mission_cooldown: dict[int, datetime] = {}
        self.ignore: set[int] = set()
        self.data_msg: dict[int, Message] = {}
        self.ocs: dict[int, Character] = {}
        self.rpers: dict[int, dict[int, Character]] = {}
        self.oc_list: dict[int, int] = {}
        self.supporting: dict[Member, Member] = {}
        self.oc_list_webhook: Optional[Webhook] = None
        self.ctx_menu1 = app_commands.ContextMenu(
            name="Read Moves",
            callback=self.moves_checker,
            guild_ids=[719343092963999804],
        )
        self.ctx_menu2 = app_commands.ContextMenu(
            name="Read Abilities",
            callback=self.abilities_checker,
            guild_ids=[719343092963999804],
        )
        self.ctx_menu3 = app_commands.ContextMenu(
            name="Check User's OCs",
            callback=self.check_ocs,
            guild_ids=[719343092963999804],
        )
        self.bot.tree.add_command(self.ctx_menu1)
        self.bot.tree.add_command(self.ctx_menu2)
        self.bot.tree.add_command(self.ctx_menu3)

    async def cog_unload(self) -> None:
        self.bot.tree.remove_command(self.ctx_menu1.name, type=self.ctx_menu1.type)
        self.bot.tree.remove_command(self.ctx_menu2.name, type=self.ctx_menu2.type)
        self.bot.tree.remove_command(self.ctx_menu3.name, type=self.ctx_menu3.type)

    async def moves_checker(self, ctx: Interaction, message: Message):
        resp: InteractionResponse = ctx.response
        await resp.defer(ephemeral=True)
        moves = []
        if oc := self.ocs.get(message.id):
            moves = list(oc.moveset.copy())
        elif text := message.content:
            moves = [
                move
                for move in Move.all()
                if move.name in text.title() or move.id in text.upper()
            ]
        if len(moves) == 1:
            view = View()
            view.add_item(
                Button(
                    label="Click here to check more information at Bulbapedia.",
                    url=moves[0].url,
                )
            )
            await ctx.followup.send(
                embed=moves[0].embed,
                ephemeral=True,
                view=view,
            )
        elif moves:
            moves.sort(key=lambda x: x.name)
            view = MoveView(
                member=ctx.user,
                moves=moves,
                target=ctx,
                keep_working=True,
            )
            async with view.send(ephemeral=True):
                self.bot.logger.info(
                    "User %s is reading the moves at %s",
                    str(ctx.user),
                    message.jump_url,
                )
        else:
            await ctx.followup.send(
                "This message does not include moves.",
                ephemeral=True,
            )

    async def abilities_checker(self, ctx: Interaction, message: Message):
        resp: InteractionResponse = ctx.response
        await resp.defer(ephemeral=True)
        abilities: list[Ability | SpAbility] = []
        if oc := self.ocs.get(message.id):
            if sp_ability := oc.sp_ability:
                abilities.append(sp_ability)
            abilities.extend(oc.abilities)
        elif text := message.content:
            abilities.extend(
                ab
                for ab in Ability.all()
                if ab.name in text.title() or ab.id in text.upper()
            )
            abilities.extend(
                sp
                for x in self.ocs.values()
                if (sp := x.sp_ability) and sp.name.lower() in text.lower()
            )

        if abilities:
            abilities.sort(key=lambda x: x.name)
            view = AbilityView(
                member=ctx.user,
                abilities=abilities,
                target=ctx,
                keep_working=True,
            )
            async with view.send(ephemeral=True):
                self.bot.logger.info(
                    "User %s is reading the abilities at %s",
                    str(ctx.user),
                    message.jump_url,
                )
        else:
            await ctx.followup.send(
                "This message does not include abilities.",
                ephemeral=True,
            )

    async def check_ocs(self, ctx: Interaction, member: Member):
        resp: InteractionResponse = ctx.response
        await resp.defer(ephemeral=True)
        ocs: list[Character] = list(self.rpers.get(member.id, {}).values())
        if len(ocs) == 1:
            view = PingView(ocs[0], ctx.user.id == ocs[0].author)
            await ctx.followup.send(
                "The user only has one character",
                ephemeral=True,
                embed=ocs[0].embed,
                view=view,
            )
        elif ocs:
            view = CharactersView(
                member=ctx.user,
                ocs=ocs,
                target=ctx,
                keep_working=True,
            )
            embed = view.embed
            embed.color = member.color
            embed.set_author(
                name=member.display_name,
                icon_url=member.display_avatar.url,
            )
            async with view.send(ephemeral=True):
                self.bot.logger.info(
                    "User %s is reading the OCs of %s",
                    str(ctx.user),
                    str(member),
                )
        else:
            await ctx.followup.send(
                f"{member.mention} has no characters.",
                ephemeral=True,
            )

    @app_commands.command(description="Grants registered role to an user")
    @app_commands.guilds(719343092963999804)
    @app_commands.checks.has_role("Moderation")
    async def register(self, ctx: Interaction, member: Member) -> None:
        """Grants registered role to an user

        Parameters
        ----------
        ctx : Context
            Context
        member : Member
            User to register
        """
        resp: InteractionResponse = ctx.response
        guild: Guild = ctx.guild
        role = guild.get_role(719642423327719434)
        author: Member = ctx.user
        await resp.defer(ephemeral=True)
        if role not in member.roles:
            await member.add_roles(role, reason=f"Registered by {author}")
            embed = Embed(
                description="You can try to use /ping `<role>` for finding a RP. "
                "(<#910914713234325504> also works)",
                colour=member.colour,
                timestamp=utcnow(),
            )
            embed.set_image(url=REGISTERED_IMG)
            embed.set_author(name=author.display_name, icon_url=author.avatar.url)
            embed.set_footer(text=guild.name, icon_url=guild.icon.url)
            files, embed = await self.bot.embed_raw(embed)

            view = View()
            view.add_item(
                Button(
                    label="Maps",
                    url="https://discord.com/channels/719343092963999804/812180282739392522/906430640898056222",
                )
            )
            view.add_item(
                Button(
                    label="Story-lines",
                    url="https://discord.com/channels/719343092963999804/903627523911458816/",
                )
            )
            view.add_item(
                Button(
                    label="RP Planning",
                    url="https://discord.com/channels/719343092963999804/722617383738540092/",
                )
            )

            with suppress(DiscordException):
                await member.send(embed=embed, files=files, view=view)
            await ctx.followup.send("User has been registered", ephemeral=True)
        else:
            await ctx.followup.send("User is already registered", ephemeral=True)

    @app_commands.command(name="ocs", description="Allows to show characters")
    @app_commands.guilds(719343092963999804)
    @app_commands.describe(member="Member, if not provided, it's current user.")
    @app_commands.describe(character="Search by name, directly")
    async def get_ocs(
        self,
        ctx: Interaction,
        member: Optional[User],
        character: Optional[CharacterArg],
    ):
        resp: InteractionResponse = ctx.response
        await resp.defer(ephemeral=True)
        if member is None:
            member = ctx.user
        if character:
            view = PingView(character, ctx.user.id == character.author)
            return await ctx.followup.send(
                embed=character.embed, view=view, ephemeral=True
            )
        if ocs := list(self.rpers.get(member.id, {}).values()):
            ocs.sort(key=lambda x: x.name)
            if len(ocs) == 1:
                await ctx.followup.send(
                    f"{member.mention} has only one character.",
                    embed=ocs[0].embed,
                    ephemeral=True,
                )
                return
            view = CharactersView(
                member=ctx.user,
                ocs=ocs,
                target=ctx,
                keep_working=True,
            )
            embed = view.embed
            embed.color = member.color
            embed.set_author(
                name=member.display_name,
                icon_url=member.display_avatar.url,
            )
            async with view.send(ephemeral=True):
                if member == ctx.user:
                    self.bot.logger.info("User %s is reading their OCs", str(member))
                else:
                    self.bot.logger.info(
                        "User %s is reading the OCs of %s",
                        str(ctx.user),
                        str(member),
                    )
        else:
            await ctx.followup.send(
                f"{member.mention} has no characters.", ephemeral=True
            )

    @app_commands.command(description="Allows to create OCs as an user")
    @app_commands.guilds(719343092963999804)
    @app_commands.describe(member="Member, if not provided, it's current user.")
    @app_commands.checks.has_role("Moderation")
    async def submit_as(self, ctx: Interaction, member: Optional[User]):
        resp: InteractionResponse = ctx.response
        await resp.defer(ephemeral=True)
        if not member:
            member = ctx.user
        if ctx.user == member:
            self.supporting.pop(ctx.user, None)
            await ctx.followup.send(
                content="OCs registered now will be assigned to your account.!",
                ephemeral=True,
            )
        else:
            self.supporting[ctx.user] = member
            await ctx.followup.send(
                content=f"OCs registered now will be assigned to {member.mention}!",
                ephemeral=True,
            )

    async def unclaiming(self, channel: Union[TextChannel, int]):
        """This method is used when a channel has been inactivate for 3 days.

        Parameters
        ----------
        channel_id : int
            channel id to use
        """
        if isinstance(channel, int):
            channel: TextChannel = self.bot.get_channel(channel)
        if not self.data_msg.get(channel.id):
            async for msg in channel.history(limit=1):
                if msg.content != CLAIM_MESSAGE:
                    msg = await channel.send(CLAIM_MESSAGE)
                self.data_msg[channel.id] = msg

        async with self.bot.database() as conn:
            for oc in filter(lambda x: x.location == channel.id, self.ocs.values()):
                await conn.execute(
                    """--sql
                    UPDATE CHARACTER
                    SET LOCATION = NULL
                    WHERE ID = $1;
                    """,
                    oc.id,
                )
                oc.location = None

    async def list_update(self, member: Object):
        """This function updates an user's character list message

        Parameters
        ----------
        member : Object
            User to update list
        """
        if not self.ready:
            return
        if isinstance(member, int):
            member = Object(id=member)
        if oc_list := self.oc_list.get(member.id):
            try:
                await self.oc_list_webhook.edit_message(oc_list, embed=None)
            except NotFound:
                oc_list = None

        if not oc_list:
            message: WebhookMessage = await self.oc_list_webhook.send(
                content=f"<@{member.id}>",
                wait=True,
                allowed_mentions=AllowedMentions(users=True),
            )
            thread = await message.create_thread(name=f"OCs⎱{member.id}")
            self.oc_list[member.id] = oc_list = thread.id
            if user := thread.guild.get_member(member.id):
                await thread.add_user(user)
            view = RPView(member.id, self.oc_list)
            await message.edit(view=view)
        return oc_list

    async def register_oc(self, oc: Type[Character]):
        member = Object(id=oc.author)
        try:
            thread_id = self.oc_list[member.id]
        except KeyError:
            thread_id = await self.list_update(member)
        oc.thread = thread_id
        guild: Guild = self.bot.get_guild(oc.server)
        user = guild.get_member(member.id) or member
        embed: Embed = oc.embed
        embed.set_image(url="attachment://image.png")
        kwargs = dict(
            content=f"<@{user.id}>",
            embed=embed,
            thread=Object(id=oc.thread),
            allowed_mentions=AllowedMentions(users=True),
        )
        if file := await self.bot.get_file(url=oc.generated_image, filename="image"):
            kwargs["file"] = file
            try:
                msg_oc = await self.oc_list_webhook.send(**kwargs, wait=True)
            except HTTPException:
                await self.list_update(member)
                kwargs["thread"] = Object(id=self.oc_list[member.id])
                msg_oc = await self.oc_list_webhook.send(**kwargs, wait=True)

            oc.id = msg_oc.id
            oc.image = msg_oc.embeds[0].image.url
            self.rpers.setdefault(user.id, {})
            self.rpers[user.id][oc.id] = oc
            self.ocs[oc.id] = oc
            self.bot.logger.info(
                "New character has been registered! > %s > %s > %s",
                str(user),
                repr(oc),
                oc.url or "Manual",
            )
            async with self.bot.database() as conn:
                await oc.update(connection=conn, idx=msg_oc.id)

    async def registration(
        self,
        ctx: Union[Interaction, Message],
        oc: Type[Character],
        worker: Member,
    ):
        """This is the function which handles the registration process,
        it will try to autocomplete data it can deduce, or ask about what
        can not be deduced.

        Parameters
        ----------
        ctx : Union[Interaction, Message]
            Message that is being interacted with
        oc : Type[Character]
            Character
        worker : Member
            User that interacts
        """

        async def send(text: str):
            if isinstance(ctx, Interaction):
                resp: InteractionResponse = ctx.response
                if not resp.is_done():
                    return await ctx.response.send_message(content=text, ephemeral=True)
                return await ctx.followup.send(content=text, ephemeral=True)
            else:
                return await ctx.reply(content=text, delete_after=5)

        if not self.ready:
            await send(
                "Bot is restarting, please be patient",
            )
            return

        if isinstance(ctx, Message):
            await send(
                "Starting submission process",
            )

        if isinstance(species := oc.species, Fakemon):  # type: ignore
            if not oc.url:
                stats_view = StatsView(
                    member=worker,
                    target=ctx,
                )
                async with stats_view:
                    if not (stats := stats_view.choice):
                        return
                    species.set_stats(*stats.value)
            if (
                sum(species.stats) > 18
                or min(species.stats) < 1
                or max(species.stats) > 5
            ):
                await send(
                    "Max stats is 18. Min 1. Max 5",
                )
                return
            if not 1 <= len(species.types) <= 2:
                view = Complex(
                    member=worker,
                    target=ctx,
                    values=Typing.all(),
                    max_values=2,
                    timeout=None,
                    parser=lambda x: (
                        str(x),
                        f"Adds the typing {x}",
                    ),
                    text_component=TextInput(
                        label="Character's Types",
                        placeholder="Type, Type",
                        required=True,
                    ),
                )
                async with view.send(
                    title="Select Typing",
                    description="Press the skip button in case you're going for single type.",
                ) as types:
                    if not types:
                        return
                    species.types = frozenset(types)
        elif isinstance(species, Fusion):  # type: ignore
            values = species.possible_types
            if not species.types:
                view = Complex(
                    member=worker,
                    target=ctx,
                    values=values,
                    max_values=1,
                    timeout=None,
                    parser=lambda x: (
                        "/".join(str(i) for i in x),
                        f"Adds the typing {'/'.join(str(i) for i in x)}",
                    ),
                    text_component=TextInput(
                        label="Fusion Typing",
                        placeholder=" | ".join(
                            "/".join(i.name for i in x).title() for x in values
                        ),
                        default="/".join(i.name for i in random_choice(values)).title(),
                    ),
                )
                async with view.send(
                    single=True,
                    title="Select Typing",
                ) as types:
                    if not types:
                        return
                    species.types = frozenset(types)
            elif oc.types not in values:
                items = ", ".join(
                    "/".join(i.name for i in item) for item in values
                ).title()
                await send(
                    f"Invalid typing for the fusion, valid types are {items}",
                )
                return

        max_ab = oc.max_amount_abilities
        if not isinstance(species, Fakemon) and (
            isinstance(species, (Legendary, Mythical, UltraBeast))
            or len(species.abilities) == 1
        ):
            oc.abilities = species.abilities
        elif not oc.abilities or len(oc.abilities) > max_ab:
            values = Ability.all() if oc.any_ability_at_first else oc.species.abilities
            placeholder = ", ".join(["Ability"] * oc.max_amount_abilities)
            ability_view = Complex(
                member=worker,
                values=values,
                target=ctx,
                max_values=max_ab,
                text_component=TextInput(
                    label="Ability",
                    style=TextStyle.paragraph,
                    placeholder=placeholder,
                    default=", ".join(
                        x.name
                        for x in sample(
                            values,
                            k=oc.max_amount_abilities,
                        )
                    ),
                ),
            )

            async with ability_view.send(
                title=f"Select the Abilities (Max {max_ab})",
                description="If you press the write button, you can add multiple by adding commas.",
            ) as abilities:
                if not abilities:
                    return
                oc.abilities = frozenset(abilities)
        if len(oc.abilities) > max_ab:
            await send(f"Max Amount of Abilities for the current Species is {max_ab}")
            return
        elif not oc.any_ability_at_first and (
            ability_errors := ", ".join(
                ability.name
                for ability in oc.abilities
                if ability not in species.abilities
            )
        ):
            await send(
                f"the abilities [{ability_errors}] were not found in the species"
            )
            return

        text_view = ModernInput(member=worker, target=ctx)

        if isinstance(species := oc.species, (Variant, Fakemon)):
            view = MovepoolView(
                target=ctx,
                member=worker,
                oc=oc,
            )
            await view.send()
            await view.wait()
            species = oc.species

        if not oc.moveset or len(oc.moveset) > 6:
            if movepool := species.total_movepool:
                movepool = movepool()
            else:
                movepool = Move.all()

            moves_view = Complex(
                member=worker,
                values=movepool,
                timeout=None,
                target=ctx,
                max_values=6,
            )

            async with moves_view.send(
                title="Select the Moves",
                description="If you press the write button, you can add multiple by adding commas.",
            ) as moves:
                if not moves:
                    return
                oc.moveset = frozenset(moves)

        if oc.url and isinstance(species, (Variant, Fakemon)):
            species.movepool += Movepool(event=oc.moveset)

        if not oc.any_move_at_first:
            moves_movepool = species.total_movepool()
            if move_errors := ", ".join(
                move.name for move in oc.moveset if move not in moves_movepool
            ):
                await send(f"the moves [{move_errors}] were not found in the movepool")
                return
        elif len(oc.moveset) > 6:
            await send("Max amount of moves in a pokemon is 6.")
            return

        if oc.sp_ability == SpAbility():
            bool_view = BooleanView(member=worker, target=ctx)
            async with bool_view.handle(
                title="Does the character have an Special Ability?",
                description=(
                    "Special abilities are basically unique traits that their OC's kind usually can't do, "
                    "it's like being born with an unique power that could have been obtained by different "
                    "reasons, they are known for having pros and cons."
                ),
            ) as answer:
                if answer is None:
                    return
                if answer:
                    data: dict[str, str] = {}
                    for item in SpAbility.__slots__:

                        if item == "name":
                            style = TextStyle.short
                        else:
                            style = TextStyle.paragraph

                        async with text_view.handle(
                            label=f"Special Ability's {item.title()}",
                            style=style,
                            placeholder=(
                                f"Here you'll define the Special Ability's {item.title()}, "
                                "make sure it is actually understandable."
                            ),
                            required=True,
                        ) as answer:
                            if not answer:
                                return
                            data[item] = answer
                    oc.sp_ability = SpAbility(**data)

        if not (oc.url or oc.backstory):
            async with text_view.handle(
                label="Character's Backstory",
                style=TextStyle.paragraph,
                placeholder=(
                    "Don't worry about having to write too much, this is just a summary of information "
                    "that people can keep in mind when interacting with your character. You can provide "
                    "information about how they are, information of their past, or anything you'd like to add."
                ),
                required=False,
            ) as text:
                if text is None:
                    return
                oc.backstory = text or None

        if not (oc.url or oc.extra):
            async with text_view.handle(
                label="Character's Extra information",
                style=TextStyle.paragraph,
                placeholder=(
                    "In this area, you can write down information you want people to consider when they are rping with them, "
                    "the information can be from either the character's height, weight, if it uses clothes, if the character likes or dislikes "
                    "or simply just writing down that your character has a goal in specific."
                ),
                required=False,
            ) as text:
                if text is None:
                    return
                oc.extra = text or None

        image_view = ImageView(
            member=worker,
            target=ctx,
            default_img=oc.image or oc.default_image,
        )
        async with image_view.send() as image:
            if image is None:
                return
            oc.image = image

        await self.register_oc(oc)

    async def oc_update(self, oc: Type[Character]):
        embed: Embed = oc.embed
        embed.set_image(url="attachment://image.png")
        try:
            try:
                await self.oc_list_webhook.edit_message(
                    oc.id,
                    embed=embed,
                    thread=Object(id=oc.thread),
                )
            except HTTPException:
                guild = self.bot.get_guild(oc.server)
                if not (thread := guild.get_thread(oc.thread)):
                    thread: Thread = await self.bot.fetch_channel(oc.thread)
                await thread.edit(archived=False)
                await self.oc_list_webhook.edit_message(
                    oc.id,
                    embed=embed,
                    thread=thread,
                )
        except NotFound:
            await self.register_oc(oc)

    async def bio_google_doc_parser(
        self, message: Message
    ) -> Optional[tuple[DocumentType, str]]:
        text: str = codeblock_converter(message.content or "").content
        if doc_data := G_DOCUMENT.match(text):
            doc = await to_thread(docs_reader, url := doc_data.group(1))
            url = f"https://docs.google.com/document/d/{url}/edit?usp=sharing"
            return doc, url

    async def bio_word_doc_parser(self, message: Message) -> Optional[DocumentType]:
        if attachments := message.attachments:
            with suppress(Exception):
                file = await attachments[0].to_file()
                return Document(file.fp)

    async def bio_discord_doc_parser(self, message: Message) -> Optional[dict]:
        text = codeblock_converter(message.content or "").content
        if G_DOCUMENT.match(text):
            return
        with suppress(MarkedYAMLError):
            text = yaml_handler(text)
            if isinstance(msg_data := safe_load(text), dict):
                if images := message.attachments:
                    msg_data["image"] = images[0].url
                return msg_data

    async def submission_handler(
        self,
        message: Interaction | Message,
        **msg_data,
    ):
        if isinstance(message, Interaction):
            refer_author = message.user
        else:
            refer_author = message.author
        if msg_data:
            author = self.supporting.get(refer_author, refer_author)
            if oc := oc_process(**msg_data):
                oc.author = author.id
                oc.server = message.guild.id
                await self.registration(ctx=message, oc=oc, worker=refer_author)
                if isinstance(message, Message):
                    with suppress(DiscordException):
                        await message.delete()

    async def on_message_submission(self, message: Message):
        """This method processes character submissions

        Attributes
        ----------
        message : Message
            Message to process
        """
        if (
            not self.ready
            or not message.guild
            or message.mentions
            or message.author.bot
            or message.author.id in self.ignore
            or message.stickers
        ):
            return
        self.ignore.add(message.author.id)
        try:
            done, _ = await wait(
                [
                    self.bio_google_doc_parser(message),
                    self.bio_discord_doc_parser(message),
                    self.bio_word_doc_parser(message),
                ],
                return_when=ALL_COMPLETED,
            )

            for result in map(lambda x: x.result(), done):
                msg_data: Optional[dict[str, Any]] = result

                if isinstance(result, tuple):
                    result, url = result
                    msg_data = doc_convert(result)
                    msg_data["url"] = url
                elif isinstance(result, DocumentType):
                    if result.tables:
                        msg_data = doc_convert(result)
                    else:
                        with suppress(MarkedYAMLError):
                            msg_data = safe_load(
                                yaml_handler(
                                    "\n".join(
                                        element
                                        for p in result.paragraphs
                                        if (element := p.text.strip())
                                    )
                                )
                            )

                if isinstance(msg_data, dict):
                    await self.submission_handler(message, **msg_data)
                    return
        except Exception as e:
            self.bot.logger.exception("Exception processing character", exc_info=e)
            await message.reply(str(e), delete_after=10)
        finally:
            self.ignore -= {message.author.id}

    async def on_message_tupper(self, message: Message, member_id: int):
        channel = message.channel
        author = message.author.name.title()

        if "Npc" in author or "Narrator" in author:
            return

        ocs = {item.name: item for item in self.rpers.get(member_id, {}).values()}

        if not (oc := ocs.get(author)):
            if items := get_close_matches(author, ocs, n=1, cutoff=0.85):
                oc = ocs[items[0]]
            elif ocs := [v for k, v in ocs.items() if k in author or author in k]:
                oc = ocs[0]
            else:
                return

        if oc.location != channel.id:
            async with self.bot.database() as db:
                oc.location = channel.id
                await oc.upsert(db)
                await self.oc_update(oc)

        if (
            (former_channel := message.guild.get_channel(oc.location))
            and not [x for x in self.ocs.values() if x.location == oc.location]
            and former_channel != channel
        ):
            await self.unclaiming(former_channel)

        if isinstance(channel, TextChannel):
            scheduler = await self.bot.scheduler.get_schedule(f"RP[{channel.id}]")
            scheduler.trigger = IntervalTrigger(days=3)

    async def on_message_proxy(self, message: Message):
        """This method processes tupper messages

        Attributes
        ----------
        message : Message
            Message to process
        """
        context = await self.bot.get_context(message)

        if context.command:
            return

        with suppress(TimeoutError):
            msg: Message = await self.bot.wait_for(
                "message",
                check=message_validator(message),
                timeout=3,
            )
            if isinstance(channel := message.channel, TextChannel):
                trigger = IntervalTrigger(days=3)
                await self.bot.scheduler.add_schedule(
                    self.unclaiming,
                    trigger,
                    id=f"RP[{channel.id}]",
                    args=[channel.id, False],
                    conflict_policy=ConflictPolicy.replace,
                )
            await self.on_message_tupper(msg, message.author.id)

    async def load_characters(self, db: Connection):
        self.bot.logger.info("Loading all Characters.")
        for oc in await fetch_all(db):
            self.ocs[oc.id] = oc
            self.rpers.setdefault(oc.author, {})
            self.rpers[oc.author][oc.id] = oc
        self.bot.logger.info("Finished loading all characters")

    async def load_profiles(self):
        self.bot.logger.info("Loading All Profiles")
        channel = await self.bot.fetch_channel(919277769735680050)
        async for m in channel.history(limit=None):
            if m.mentions and m.webhook_id:
                user = m.mentions[0]
                self.oc_list[user.id] = m.id
                view = RPView(user.id, self.oc_list)
                self.bot.add_view(view=view, message_id=m.id)

        self.bot.logger.info("Finished loading all Profiles.")

    async def load_missions(self, db: Connection):
        self.bot.logger.info("Loading claimed missions")
        missions: dict[int, Mission] = {}

        async for item in db.cursor("SELECT * FROM MISSIONS;"):
            mission = Mission(**dict(item))
            if mission.id:
                self.missions.add(mission)
                missions[mission.id] = mission

        async for oc_item in db.cursor("SELECT * FROM MISSION_ASSIGNMENT;"):
            mission_id, oc_id, assigned_at = (
                oc_item["mission"],
                oc_item["character"],
                oc_item["assigned_at"],
            )
            if (mission := missions.get(mission_id)) and (oc := self.ocs.get(oc_id)):
                mission.ocs |= {oc.id}
                self.mission_claimers.setdefault(mission.id, set())
                self.mission_claimers[mission.id].add(oc.id)
                self.mission_cooldown.setdefault(oc.author, assigned_at)
                if self.mission_cooldown[oc.author] < assigned_at:
                    self.mission_cooldown[oc.author] = assigned_at

        self.bot.logger.info("Finished loading claimed missions")

    async def load_mission_views(self, db: Connection):
        self.bot.logger.info("Loading mission views")

        channel: TextChannel = await self.bot.fetch_channel(908498210211909642)

        for mission in self.missions:
            view = MissionView(
                mission=mission,
                mission_claimers=self.mission_claimers,
                mission_cooldown=self.mission_cooldown,
                supporting=self.supporting,
            )
            try:
                message = await channel.fetch_message(mission.msg_id)
                await message.edit(view=view)
            except DiscordException:
                if not (member := channel.guild.get_member(mission.author)):
                    return await mission.remove(db)
                msg = await channel.send(
                    content=member.mention,
                    embed=mission.embed,
                    view=view,
                    allowed_mentions=AllowedMentions(users=True),
                )
                mission.msg_id = msg.id
                thread = await msg.create_thread(name=f"Mission {mission.id:03d}")
                await thread.add_user(member)
                ocs = set(mission.ocs)
                for oc_id in mission.ocs:
                    if oc := self.ocs.get(oc_id):
                        view = View()
                        view.add_item(Button(label="Jump URL", url=oc.jump_url))
                        await thread.send(
                            f"{member} joined with {oc.name} `{oc!r}` as character for this mission.",
                            view=view,
                        )
                    else:
                        ocs.remove(oc_id)
                mission.ocs = frozenset(ocs)

                mission.msg_id = msg.id
                await mission.upsert(db)

        self.bot.logger.info("Finished loading mission views")

    async def load_submssions(self):
        self.bot.logger.info("Loading Submission menu")
        source = Path("resources/templates.json")
        async with aiopen(source.resolve(), mode="r") as f:
            contents = await f.read()
            view = SubmissionView(
                ocs=self.ocs,
                rpers=self.rpers,
                oc_list=self.oc_list,
                supporting=self.supporting,
                missions=self.missions,
                mission_claimers=self.mission_claimers,
                mission_cooldown=self.mission_cooldown,
                **loads(contents),
            )
            webhook = await self.bot.webhook(852180971985043466)
            await webhook.edit_message(
                961345742222536744,
                content=None,
                view=view,
            )
        self.bot.logger.info("Finished loading Submission menu")

    async def load_claimed_categories(self):
        items: list[list[TextChannel]] = [
            [
                x
                for x in self.bot.get_channel(ch).channels
                if "\N{RIGHT-POINTING DOUBLE ANGLE QUOTATION MARK}" not in x.name
            ]
            for ch in RP_CATEGORIES
        ]
        for channel in chain(*items):
            async for m in channel.history(limit=1):

                date = m.created_at + timedelta(days=3)

                if m.author == self.bot.user and m.content == CLAIM_MESSAGE:
                    self.data_msg[channel.id] = m

                trigger = IntervalTrigger(days=3, start_time=date)

                await self.bot.scheduler.add_schedule(
                    self.unclaiming,
                    trigger=trigger,
                    id=f"RP[{channel.id}]",
                    args=[channel.id],
                    conflict_policy=ConflictPolicy.replace,
                )

    @commands.Cog.listener()
    async def on_message(self, message: Message) -> None:
        """on_message handler

        Parameters
        ----------
        message : Message
            Message to process
        """
        if message.channel.id == 852180971985043466:
            await self.on_message_submission(message)
        elif (
            message.guild
            and (tupper := message.guild.get_member(431544605209788416))
            and tupper.status == Status.online
            and message.channel.category_id in RP_CATEGORIES
            and not message.webhook_id
            and "\N{RIGHT-POINTING DOUBLE ANGLE QUOTATION MARK}"
            not in message.channel.name
        ):
            await self.on_message_proxy(message)

    @commands.Cog.listener()
    async def on_message_edit(self, _: Message, message: Message):
        """on_message_edit handler

        Parameters
        ----------
        _ : Message
            Previous message
        message : Message
            Message to process
        """
        if message.channel.id == 852180971985043466:
            await self.on_message_submission(message)

    @commands.Cog.listener()
    async def on_thread_delete(self, thread: Thread) -> None:
        """Detects if threads were removed

        Parameters
        ----------
        payload : RawThreadDeleteEvent
            Information
        """
        if thread.parent_id != 919277769735680050:
            return
        ocs = [oc for oc in self.ocs.values() if oc.thread == thread.id]

        if not ocs:
            return

        self.oc_list.pop(ocs[0].author, None)
        async with self.bot.database() as db:
            for oc in ocs:
                self.ocs.pop(oc.id, None)
                self.bot.logger.info(
                    "Character Removed as Thread was removed! > %s - %s > %s",
                    oc.name,
                    repr(oc),
                    oc.url or "None",
                )
                await oc.delete(db)

    @commands.Cog.listener()
    async def on_raw_message_delete(self, payload: RawMessageDeleteEvent) -> None:
        """Detects if ocs or lists were deleted

        Parameters
        ----------
        payload : RawMessageDeleteEvent
            Information
        """
        if oc := self.ocs.get(payload.message_id):
            del self.ocs[oc.id]
            self.rpers.setdefault(oc.author, {})
            self.rpers[oc.author].pop(oc.id, None)
            async with self.bot.database() as db:
                self.bot.logger.info(
                    "Character Removed as message was removed! > %s - %s > %s",
                    oc.name,
                    repr(oc),
                    oc.url or "None",
                )
                await oc.delete(db)
        if payload.message_id in self.oc_list.values():
            author_id: int = [
                k for k, v in self.oc_list.items() if v == payload.message_id
            ][0]
            del self.oc_list[author_id]
            async with self.bot.database() as db:
                for oc in self.rpers.pop(author_id, {}).values():
                    del self.ocs[oc.id]
                    self.bot.logger.info(
                        "Character Removed as Thread was removed! > %s > %s",
                        str(type(oc)),
                        oc.url or "None",
                    )
                    await oc.delete(db)

    @commands.Cog.listener()
    async def on_ready(self) -> None:
        """On ready, the parameters from Cog submisisons are loaded."""

        if self.ready:
            return

        self.oc_list_webhook = await self.bot.webhook(919277769735680050)
        async with self.bot.database() as db:
            await self.load_characters(db)
            await self.load_missions(db)
            await self.load_mission_views(db)

        await self.load_profiles()
        await self.load_submssions()
        await self.load_claimed_categories()
        self.ready = True


async def setup(bot: CustomBot) -> None:
    """Default Cog loader

    Parameters
    ----------
    bot: CustomBot
        Bot
    """
    await bot.add_cog(Submission(bot))
