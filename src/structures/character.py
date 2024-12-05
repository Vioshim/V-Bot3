# Copyright 2022 Vioshim
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#      https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.


from __future__ import annotations

import math
import re
from dataclasses import asdict, dataclass, field
from enum import Enum, StrEnum
from io import BytesIO
from typing import Any, Iterable, Optional, Type

from discord import Color, Embed, File, Interaction
from discord.app_commands import Choice
from discord.app_commands.transformers import Transform, Transformer
from discord.utils import snowflake_time, utcnow
from docx import Document as document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from motor.motor_asyncio import AsyncIOMotorCollection
from rapidfuzz import process
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Flowable,
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.structures.ability import SpAbility
from src.structures.bot import CustomBot
from src.structures.mon_typing import TypingEnum
from src.structures.move import Move
from src.structures.movepool import Movepool
from src.structures.pokeball import Pokeball
from src.structures.pronouns import Pronoun
from src.structures.species import (
    CustomSpecies,
    Fakemon,
    Fusion,
    Legendary,
    Mega,
    Mythical,
    Paradox,
    Pokemon,
    Species,
    UltraBeast,
    Variant,
)
from src.utils.functions import common_pop_get, fix, int_check
from src.utils.imagekit import Fonts, ImageKit
from src.utils.matches import G_DOCUMENT

__all__ = ("Character", "CharacterArg", "Kind", "Size")


class AgeGroup(Enum):
    Unknown = 0, "The age is Unknown.", 1.0
    Egg = 1, "Newly hatched Pokémon.", 0.1
    Baby = 5, "Recently evolved from an egg, still in early stages", 0.3
    Toddler = 10, "Learning and developing, but not quite a child.", 0.5
    Child = 15, "Considered a child in their species.", 0.75
    Adolescent = 18, "15 - 17 (Rough equivalent in human years)", 0.85
    Young_Adult = 26, "18 - 25 (Rough equivalent in human years)", 1.0
    Adult = 40, "26 - 39 (Rough equivalent in human years)", 1.0
    Middle_Aged = 55, "40 - 54 (Rough equivalent in human years)", 0.95
    Senior = 70, "55 - 69 (Rough equivalent in human years)", 0.90
    Elderly = 85, "70 - 84 (Rough equivalent in human years)", 0.85
    Ancient = 100, "85 - 99 (Rough equivalent in human years)", 0.80
    Timeless = 125, "Magic, pokeball from past, or lived long enough.", 1.0

    @classmethod
    def parse(cls, item: AgeGroup | Optional[int] | str):
        if isinstance(item, AgeGroup):
            return item

        item = item or 0
        if isinstance(item, str):
            if item.isdigit():
                item = int(item)
            elif foo := process.extractOne(
                item,
                cls,
                processor=lambda x: x.name if isinstance(x, cls) else x,
                score_cutoff=85,
            ):
                return foo[0]

        return cls.from_number(item) if isinstance(item, int) else cls.Unknown

    @classmethod
    def from_number(cls, item: Optional[int]):
        base = item or 0
        return next((x for x in cls if x.value[0] >= base), cls.Ancient)

    @property
    def title(self):
        return self.name.replace("_", " ")

    @property
    def key(self):
        value, _, _ = self.value
        return value

    @property
    def description(self):
        _, desc, _ = self.value
        return desc

    @property
    def scale(self):
        _, _, scale = self.value
        return scale


class Kind(Enum):
    Common = Pokemon
    Legendary = Legendary
    Mythical = Mythical
    UltraBeast = UltraBeast
    Fakemon = Fakemon
    Variant = Variant
    Mega = Mega
    Fusion = Fusion
    Paradox = Paradox

    @property
    def title(self):
        name = self.name
        if name.startswith("Custom"):
            name = name.removeprefix("Custom") + " (Custom)"
        return name.replace("Common", "Pokemon")

    @property
    def to_db(self):
        match self:
            case self.Common:
                return "COMMON"
            case self.Legendary:
                return "LEGENDARY"
            case self.Mythical:
                return "MYTHICAL"
            case self.UltraBeast:
                return "ULTRA BEAST"
            case self.Fakemon:
                return "FAKEMON"
            case self.Variant:
                return "VARIANT"
            case self.Paradox:
                return "PARADOX"
            case self.Mega:
                return "MEGA"
            case self.Fusion:
                return "FUSION"

    @classmethod
    def associated(cls, name: str) -> Optional[Kind]:
        match fix(name):
            case "COMMON":
                return cls.Common
            case "LEGENDARY":
                return cls.Legendary
            case "MYTHICAL":
                return cls.Mythical
            case "ULTRABEAST":
                return cls.UltraBeast
            case "FAKEMON":
                return cls.Fakemon
            case "VARIANT":
                return cls.Variant
            case "PARADOX":
                return cls.Paradox
            case "MEGA":
                return cls.Mega
            case "FUSION":
                return cls.Fusion

    def all(self) -> frozenset[Species]:
        return self.value.all()

    @classmethod
    def from_class(cls, item: Type[Species]):
        return next(
            (element for element in cls if isinstance(item, element.value)),
            cls.Fakemon,
        )


class Size(float, Enum):
    Titan = 2.1
    Maximum = 1.5
    Very_Large = 1.375
    Large = 1.25
    Above_Average = 1.125
    Average = 1.0
    Below_Average = 0.875
    Small = 0.75
    Very_Small = 0.625
    Minimum = 0.5
    Minuscule = 0.25

    @property
    def emoji(self):
        if self.value > 1:
            return "🟧"
        if self.value < 1:
            return "🟦"
        return "🟩"

    @property
    def reference_name(self):
        return "Default" if self == Size.Average else self.name.replace("_", " ")

    def height_value(self, value: float = 0):
        if value:
            return round(value * self, 4)
        return round(0.825 * self.value**2, 4)

    def weight_value(self, value: float = 0):
        if value:
            return round(value * self.value, 4)
        return round(10 * ((self.value - 0.25) / 0.6) ** 3, 4)

    @staticmethod
    def meters_to_ft_inches(value: float = 0):
        return int(value / 0.3048), float(value / 0.3048 % 1 * 12)

    @staticmethod
    def kg_to_lbs(value: float = 0):
        return math.ceil(value * 220.462) / 100

    @staticmethod
    def ft_inches_to_meters(feet: float = 0, inches: float = 0):
        return math.ceil(feet * 30.48 + inches * 2.54) / 100

    @staticmethod
    def lbs_to_kgs(value: float = 0):
        return math.ceil(value * 45.359) / 100

    def height_info(self, value: float = 0):
        value = self.height_value(value)

        if value >= 5.87862537e10:
            return f"{value/9.461e+15:.2f} ly / {value/5.87862537e12:.2f} au"

        if value >= 1000:
            return f"{value/1000:.2f} km / {value/1609.34:.2f} mi"

        value_ft, value_in = value // 0.3048, value / 0.3048 % 1 * 12

        if value_ft > 0:
            return f"{value:.2f} m / {value_ft:.0f}' {value_in:.0f}\" ft"

        if value_in >= 0.1:
            return f"{value*100:.2f} cm / {value_in:.2f} in"

        return f"{value*1000:.2f} mm / {value_in*1000:.2f} th"

    def weight_info(self, value: float = 0):
        value = self.weight_value(value)
        lbs = self.kg_to_lbs(value)
        if value < 1:
            return f"{value*1000:.2f} g / {lbs*16:.2f} oz"
        return f"{value:.2f} kg / {lbs:.2f} lbs"


class Stats(StrEnum):
    HP = "HP"
    ATK = "Attack"
    DEF = "Defense"
    SPA = "Sp. Attack"
    SPD = "Sp. Defense"
    SPE = "Speed"


class Nature(Enum):
    def __init__(self, high: Stats, low: Stats):
        self.description = f"⬆ {high.value} | ⬇ {low.value}"

    # fmt: off
    Composed   = (Stats.HP,  Stats.HP )  # noqa: E202,E221
    Cuddly     = (Stats.HP,  Stats.ATK)  # noqa: E221
    Distracted = (Stats.HP,  Stats.DEF)  # noqa: E221
    Proud      = (Stats.HP,  Stats.SPA)  # noqa: E221
    Decisive   = (Stats.HP,  Stats.SPD)  # noqa: E221
    Patient    = (Stats.HP,  Stats.SPE)  # noqa: E221
    Desperate  = (Stats.ATK, Stats.HP )  # noqa: E202,E221
    Hardy      = (Stats.ATK, Stats.ATK)  # noqa: E221
    Lonely     = (Stats.ATK, Stats.DEF)  # noqa: E221
    Adamant    = (Stats.ATK, Stats.SPA)  # noqa: E221
    Naughty    = (Stats.ATK, Stats.SPD)  # noqa: E221
    Brave      = (Stats.ATK, Stats.SPE)  # noqa: E221
    Stark      = (Stats.DEF, Stats.HP )  # noqa: E202,E221
    Bold       = (Stats.DEF, Stats.ATK)  # noqa: E221
    Docile     = (Stats.DEF, Stats.DEF)  # noqa: E221
    Impish     = (Stats.DEF, Stats.SPA)  # noqa: E221
    Lax        = (Stats.DEF, Stats.SPD)  # noqa: E221
    Relaxed    = (Stats.DEF, Stats.SPE)  # noqa: E221
    Curious    = (Stats.SPA, Stats.HP )  # noqa: E202,E221
    Modest     = (Stats.SPA, Stats.ATK)  # noqa: E221
    Mild       = (Stats.SPA, Stats.DEF)  # noqa: E221
    Bashful    = (Stats.SPA, Stats.SPA)  # noqa: E221
    Rash       = (Stats.SPA, Stats.SPD)  # noqa: E221
    Quiet      = (Stats.SPA, Stats.SPE)  # noqa: E221
    Dreamy     = (Stats.SPD, Stats.HP )  # noqa: E202,E221
    Calm       = (Stats.SPD, Stats.ATK)  # noqa: E221
    Gentle     = (Stats.SPD, Stats.DEF)  # noqa: E221
    Careful    = (Stats.SPD, Stats.SPA)  # noqa: E221
    Quirky     = (Stats.SPD, Stats.SPD)  # noqa: E221
    Sassy      = (Stats.SPD, Stats.SPE)  # noqa: E221
    Skittish   = (Stats.SPE, Stats.HP )  # noqa: E202,E221
    Timid      = (Stats.SPE, Stats.ATK)  # noqa: E221
    Hasty      = (Stats.SPE, Stats.DEF)  # noqa: E221
    Jolly      = (Stats.SPE, Stats.SPA)  # noqa: E221
    Naive      = (Stats.SPE, Stats.SPD)  # noqa: E221
    Serious    = (Stats.SPE, Stats.SPE)  # noqa: E221
    # fmt: on


class Gender(Enum):
    Male = "Has male features in its body."
    Female = "Has female features in its body."
    Genderless = "Lacks distinct male or female features."
    Fluid = "Uses transformations to change gender frequently."


class SizeKind(float, Enum):
    Kaiju = 10.0
    Regular = 1.0
    Shrunk = 0.1

    @property
    def emoji(self):
        match self:
            case self.Kaiju:
                return "🟧"
            case self.Regular:
                return "🟩"
            case self.Shrunk:
                return "🟦"

    def is_valid(self):
        return 0.1 <= self.value <= 10.0

    def label_for(self, scale: float = 1.0, minimum: float = 1.0, maximum: float = 2.0, real: bool = False):
        n = self.name.replace("_", " ")
        ma = maximum * scale
        mi = minimum * scale

        if real:
            ma *= self.value
            mi *= self.value

        if mi >= 5.87862537e10:
            return f"{n} ({mi/9.461e+15:.2f} ly / {mi/5.87862537e12:.2f} au - {ma/9.461e+15:.2f} ly / {ma/5.87862537e12:.2f} au)"

        if ma >= 1000:
            return f"{n} ({mi/1000:.2f} km / {mi/1609.34:.2f} mi - {ma/1000:.2f} km / {ma/1609.34:.2f} mi)"

        ma_ft, ma_in = int(ma / 0.3048), ma / 0.3048 % 1 * 12
        mi_ft, mi_in = int(mi / 0.3048), mi / 0.3048 % 1 * 12

        if ma_ft >= 100:
            return f"{n} ({mi:.2f} m / {mi/0.9144:.2f} yd - {ma:.2f} m / {ma/0.9144:.2f} yd)"

        if ma_ft >= 1:
            return f"{n} ({mi:.2f} m / {mi_ft:02d}' {mi_in:.0f}\" ft - {ma:.2f} m / {ma_ft:02d}' {ma_in:.0f}\" ft)"

        if mi_in >= 0.1:
            return f"{n} ({mi*100:.2f} cm / {mi_in:.2f} in - {ma*100:.2f} cm / {ma_in:.2f} in)"

        return f"{n} ({mi*1000:.2f} mm / {mi_in*1000:.2f} th - {ma*1000:.2f} mm / {ma_in*1000:.2f} th)"


class Weight(float, Enum):
    Very_Robust = 1.500
    Obese = Very_Robust

    Slightly_Robust = 1.375
    Very_Overweight = Slightly_Robust

    Very_Sturdy = 1.250
    Overweight = Very_Sturdy

    Slightly_Sturdy = 1.125
    Slightly_Overweight = Slightly_Sturdy

    Balanced = 1.000
    Average = Balanced

    Slightly_Lean = 0.875
    Slightly_Underweight = Slightly_Lean

    Very_Lean = 0.750
    Underweight = Very_Lean

    Slightly_Delicate = 0.625
    Very_Underweight = Slightly_Delicate

    Very_Delicate = 0.500
    Scrawny = Very_Delicate


@dataclass(slots=True)
class Character:
    species: Optional[Species] = None
    id: int = 0
    author: Optional[int] = None
    thread: Optional[int] = None
    server: int = 1196879060173852702
    name: str = ""
    age: AgeGroup = AgeGroup.Unknown
    pronoun: frozenset[Pronoun] = field(default_factory=frozenset)
    backstory: Optional[str] = None
    personality: Optional[str] = None
    extra: Optional[str] = None
    moveset: frozenset[Move] = field(default_factory=frozenset)
    sp_ability: Optional[SpAbility] = None
    url: str = ""
    image: Optional[int | str | File] = None
    location: Optional[int] = None
    hidden_power: Optional[TypingEnum] = None
    size: float = 1.65
    size_kind: SizeKind = SizeKind.Regular
    weight: Weight = Weight.Average
    last_used: Optional[int] = None
    nature: Optional[Nature] = None
    hidden_info: Optional[str] = None
    pokeball: Optional[Pokeball] = None
    gender: Gender = Gender.Genderless
    static: bool = False
    template: Optional[str] = None

    @classmethod
    def from_dict(cls, kwargs: dict[str, Any]) -> Character:
        kwargs = {k.lower(): v for k, v in kwargs.items() if k.lower() in cls.__slots__}
        return Character(**kwargs)

    def to_mongo_dict(self):
        data = asdict(self)
        data["pokeball"] = self.pokeball and self.pokeball.name
        data["species"] = self.species and self.species.as_data()
        data["age"] = self.age.name
        data["size"] = self.size
        data["weight"] = self.weight.name
        data["pronoun"] = [x.name for x in self.pronoun]
        data["moveset"] = [x.id for x in self.moveset]
        data["hidden_power"] = self.hidden_power.name if self.hidden_power else None
        data["nature"] = self.nature.name if self.nature else None
        data["gender"] = self.gender and self.gender.name
        if isinstance(self.sp_ability, SpAbility):
            aux = asdict(self.sp_ability)
            aux["kind"] = self.sp_ability.kind.name
            data["sp_ability"] = aux
        if isinstance(self.image, File):
            data["image"] = None
        return data

    @classmethod
    def from_mongo_dict(cls, dct: dict[str, Any]):
        dct.pop("_id", None)
        species = dct.pop("species", None)
        if trope := dct.get("trope", []):
            dct["tropes"] = trope
        dct["species"] = species and Species.from_data(species)
        dct.pop("size_category", None)
        return Character.from_dict(dct)

    def copy(self):
        data = self.to_mongo_dict()
        return Character.from_mongo_dict(data)

    def __post_init__(self):
        self.image_url = self.image
        self.url = self.url or ""
        if isinstance(self.species, str):
            self.species = Species.from_ID(self.species)
        if not self.server:
            self.server = 1196879060173852702
        if isinstance(self.sp_ability, dict):
            self.sp_ability = SpAbility(**self.sp_ability)
        if isinstance(self.moveset, str):
            self.moveset = [self.moveset]
        self.moveset = Move.deduce_many(*self.moveset)

        try:
            self.size = Size[self.size].height_value(1.65) if isinstance(self.size, str) else float(self.size)
        except (KeyError, ValueError):
            self.size = 1.65

        if not isinstance(self.gender, Gender):
            try:
                self.gender = Gender[self.gender]
            except KeyError:
                self.gender = Gender.Genderless

        if not isinstance(self.weight, Weight):
            try:
                self.weight = Weight[self.weight]
            except KeyError:
                self.weight = Weight.Average

        if isinstance(self.pokeball, str):
            try:
                self.pokeball = Pokeball[self.pokeball]
            except KeyError:
                self.pokeball = None

        if isinstance(self.nature, str):
            try:
                self.nature = Nature[self.nature]
            except KeyError:
                self.nature = None

        if isinstance(self.pronoun, str):
            self.pronoun = [self.pronoun]
        self.pronoun = Pronoun.deduce_many(*self.pronoun)

        self.age = AgeGroup.parse(self.age)
        if self.hidden_power:
            self.hidden_power = TypingEnum.deduce(self.hidden_power)

        if isinstance(self.size_kind, float):
            try:
                self.size_kind = SizeKind(self.size_kind)
            except ValueError:
                self.size_kind = SizeKind.Regular

        self.size = min(2.1 * self.age.scale, max(1.2 * self.age.scale, self.size))

    def __eq__(self, other: Character):
        return isinstance(other, Character) and self.id == other.id

    def __ne__(self, other: Character) -> bool:
        return isinstance(other, Character) and other.id != self.id

    def __hash__(self) -> int:
        return self.id >> 22

    @property
    def min_amount_species(self):
        return 2 if isinstance(self.species, Fusion) else 1

    @property
    def max_amount_species(self):
        return 3 if isinstance(self.species, Fusion) else 1

    @property
    def last_used_at(self):
        data = max(self.id or 0, self.location or 0, self.last_used or 0)
        return snowflake_time(data) if data else utcnow()

    @property
    def types(self) -> frozenset[TypingEnum]:
        mon_types = [TypingEnum.Typeless]
        if self.species and self.species.types:
            mon_types = self.species.types
        return frozenset(mon_types)

    @property
    def possible_types(self) -> frozenset[frozenset[TypingEnum]]:
        return self.species.possible_types if self.species else frozenset()

    @property
    def kind(self):
        return Kind.from_class(self.species)

    @property
    def image_url(self):
        if isinstance(self.image, int) and self.thread:
            return f"https://media.discordapp.net/attachments/{self.thread}/{self.image}/image.png"
        if isinstance(self.image, str) and self.image.startswith(
            "https://media.discordapp.net/attachments/1196967811034132581/"
        ):
            return self.image

    @image_url.setter
    def image_url(self, url: str):
        if isinstance(url, str) and self.thread:
            if find := re.match(
                rf"https:\/\/\w+\.discordapp\.\w+\/attachments\/{self.thread}\/(\d+)\/image\.png",
                string=url,
            ):
                url = int(find[1])

        self.image = url or None

    @image_url.deleter
    def image_url(self):
        self.image = None

    @property
    def created_at(self):
        return snowflake_time(self.id) if self.id else utcnow()

    @property
    def document_url(self):
        return f"https://docs.google.com/document/d/{self.url}/edit?usp=sharing" if self.url else ""

    @document_url.setter
    def document_url(self, url: str):
        if item := G_DOCUMENT.match(url):
            url = item.group(1)
        else:
            url = url.removeprefix("https://docs.google.com/document/d/")
            url = url.removesuffix("/edit?usp=sharing")
        self.url = url

    @document_url.deleter
    def document_url(self):
        self.url = ""

    @property
    def evolves_from(self):
        return self.species and self.species.species_evolves_from

    @property
    def evolves_to(self):
        return self.species and self.species.species_evolves_to

    @property
    def pronoun_emoji(self):
        match self.pronoun:
            case x if Pronoun.He in x and Pronoun.She not in x:
                return Pronoun.He.emoji
            case x if Pronoun.She in x and Pronoun.He not in x:
                return Pronoun.She.emoji
            case _:
                return Pronoun.Them.emoji

    @property
    def default_emoji(self):
        return self.pronoun_emoji

    @property
    def pronoun_text(self):
        return ", ".join(sorted(x.name for x in self.pronoun)) or "Them"

    @property
    def movepool(self) -> Movepool:
        return self.species.movepool if self.species else Movepool()

    @property
    def total_movepool(self) -> Movepool:
        movepool = Movepool()
        if self.species:
            movepool += self.species.total_movepool
        return movepool

    @property
    def place_mention(self):
        if self.location:
            return f"<#{self.location}>"

    @property
    def jump_url(self):
        if self.server and self.thread and self.id:
            return f"https://discord.com/channels/{self.server}/{self.thread}/{self.id}"

    @property
    def location_url(self):
        url = f"https://discord.com/channels/{self.server}/"
        url += f"{self.location}" if self.location else ""
        return f"{url}/{self.last_used}" if self.last_used else url

    @property
    def height_text(self):
        return Size.Average.height_info(self.height_value)

    @property
    def real_height_text(self):
        return Size.Average.height_info(self.real_height_value)

    @property
    def height_value(self):
        if isinstance(self.size, Size):
            value = self.size.height_value(1.65)
        else:
            value = self.size
        return round(value, 4)

    @property
    def real_height_value(self):
        return round(self.height_value * self.size_kind.value, 4)

    @property
    def weight_text(self):
        return self.weight.name.replace("_", " ")

    @property
    def weight_value(self):
        if isinstance(self.weight, Size):
            value = self.weight.weight_value(60)
        else:
            value = self.weight
        return round(value, 4)

    @property
    def default_image(self):
        """This allows to obtain a default image for the character

        Returns
        -------
        Optional[str]
            Image it defaults to
        """
        return self.species and self.species.base_image

    @property
    def species_data(self):
        match self.species:
            case mon if isinstance(mon, Fusion):
                mon_name = "\n".join(f"* {x.name}" for x in mon.bases)[:1024]
            case mon if isinstance(mon, Fakemon):
                if (evolves_from := mon.species_evolves_from) and evolves_from.name != mon.name:
                    return f"{evolves_from.name} Evo", mon.name
                mon_name = mon.name
            case mon if isinstance(mon, CustomSpecies) and mon.base:
                mon_name = mon.name or mon.base.name
            case mon if isinstance(mon, Species):
                mon_name = mon.name
            case _:
                return "Unknown", "Unknown"

        return "Species", mon_name

    @property
    def embed(self) -> Embed:
        return self.embeds[0]

    @property
    def color(self) -> Color:
        if hidden_power := self.hidden_power:
            return Color(hidden_power.color)
        return Color.blurple()

    @property
    def embeds(self) -> list[Embed]:
        """Discord embed out of the character

        Returns
        -------
        list[Embed]
            Embed with the character's information
        """
        c_embed = Embed(title=self.name[:256].title(), url=self.document_url, timestamp=self.created_at)
        embeds = [c_embed]

        if backstory := self.backstory:
            c_embed.description = backstory[:2000]

        gender_text = self.gender.name if self.gender != Gender.Genderless else "Pronouns"
        c_embed.add_field(name=gender_text, value=self.pronoun_text or "Unknown")
        c_embed.add_field(name="Age", value=self.age.title)

        if species_data := self.species_data:
            name1, name2 = species_data
            c_embed.add_field(name=name1, value=name2)

        if (sp_ability := self.sp_ability) and sp_ability.valid:
            sp_embed = Embed(
                title=name if (name := sp_ability.name[:100]) else f"{self.name[:92]}'s Trait",
                description=sp_ability.description[:1024],
                timestamp=self.created_at,
            )

            if origin := sp_ability.origin[:600]:
                sp_embed.add_field(name="Origin", value=origin, inline=False)

            if pros := sp_ability.pros[:600]:
                sp_embed.add_field(name="Pros", value=pros, inline=False)

            if cons := sp_ability.cons[:600]:
                sp_embed.add_field(name="Cons", value=cons, inline=False)

            sp_embed.set_footer(text=sp_ability.kind.phrase)

            embeds.append(sp_embed)

        phrase = "OC's Moveset" if not self.static else "NPC's Moveset"
        if hidden_power := self.hidden_power:
            color = Color(hidden_power.color)
            moveset_title = f"{phrase} - H. Power: {hidden_power.name}"
        else:
            color = Color.blurple()
            moveset_title = phrase

        if self.pokeball:
            embeds[-1].set_thumbnail(url=self.pokeball.url)

        embeds[0].color, embeds[-1].color = color, color
        footer_elements: list[str] = []

        if self.nature:
            footer_elements.append(f"Nature: {self.nature.name}")
        footer_elements.append(f"Types: {', '.join(x.name for x in self.types)}")
        ref = self.template or self.size_kind.name
        footer_elements.append(f"{ref}: {self.height_text}")
        footer_elements.append(f"Shape: {self.weight_text}")
        footer_text = "\n".join(footer_elements) or "No additional information."
        c_embed.set_footer(text=footer_text)

        def move_parser(x: Move):
            item = self.hidden_power if x.move_id in {237, 851} and self.hidden_power else x.type
            item = TypingEnum.Typeless if TypingEnum.Typeless in self.types else item
            return f"* [{x.name}] - {item.name} ({x.category.name})".title()

        moves_text = "\n".join(map(move_parser, sorted(self.moveset, key=lambda x: x.name)))
        c_embed.add_field(name=moveset_title, value=moves_text or "> No information.", inline=False)

        if image := self.image_url:
            c_embed.set_image(url=image)
        elif isinstance(self.image, File):
            c_embed.set_image(url=f"attachment://{self.image.filename}")
        elif isinstance(self.image, str) and self.image:
            c_embed.set_image(url=self.image)
        else:
            c_embed.set_image(url="attachment://image.png")

        if self.personality:
            c_embed.add_field(name="Personality", value=self.personality[:512], inline=False)

        if self.extra:
            c_embed.add_field(name="Extra", value=self.extra[:512], inline=False)

        return embeds

    @property
    def params_header(self):
        data = {
            "Age": self.age.title,
            "Hidden Power": self.hidden_power.name if self.hidden_power else "Unknown",
            "Nature": self.nature.name if self.nature else None,
        }

        data["Measure"] = "\n".join([*self.height_text.split(" / "), f"Weight: {self.weight_text}"])

        if species_data := self.species_data:
            data[species_data[0]] = species_data[1]

        return data

    async def to_docx(self, bot: CustomBot):
        doc: Document = document()

        params_header = self.params_header

        doc.add_heading(f"{self.pronoun_emoji}〛{self.name}", 0)

        table = doc.add_table(rows=1, cols=len(params_header))
        hdr_cells = table.rows[0].cells
        for text, row in zip(params_header, hdr_cells):
            row.text = text

        row_cells = table.add_row().cells
        for index, item in enumerate(params_header.values()):
            row_cells[index].text = str(item)

        if img_file := await bot.get_file(self.image_url):
            doc.add_picture(img_file.fp, width=Inches(6))

        if self.moveset:
            doc.add_heading("Favorite Moves", level=1)
            items = sorted(self.moveset, key=lambda x: x.name)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            for index, values in enumerate((items[:2], items[2:])):
                move_args = []

                for item in values:
                    item_type = item.type
                    if item.name in ["Hidden Power", "Tera Blast"] and self.hidden_power:
                        item_type = self.hidden_power
                    item_type = TypingEnum.Typeless if TypingEnum.Typeless in self.types else item_type
                    move_args.append(f"• {item.name} - {item.category.name} - {item_type.name}".title())

                hdr_cells[index].text = "\n".join(move_args)

        if self.backstory or self.extra or self.personality:
            doc.add_page_break()

        if self.backstory:
            doc.add_heading("Bio", 1)
            doc.add_paragraph(self.backstory).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if self.personality:
            doc.add_heading("Personality", 1)
            doc.add_paragraph(self.personality).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if self.extra:
            doc.add_heading("Extra Information", 1)
            doc.add_paragraph(self.extra).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if sp_ability := self.sp_ability:
            doc.add_page_break()
            doc.add_heading(f"Unique Trait: {sp_ability.name}", 1)
            doc.add_heading(sp_ability.kind.phrase, 2)
            if sp_ability.description:
                doc.add_paragraph(sp_ability.description).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if origin := sp_ability.origin:
                doc.add_heading("How was it obtained?", 2)
                doc.add_paragraph(origin).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if pros := sp_ability.pros:
                doc.add_heading("How does it make the character's life easier?", 2)
                doc.add_paragraph(pros).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if cons := sp_ability.cons:
                doc.add_heading("How does it make the character's life harder?", 2)
                doc.add_paragraph(cons).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if (
            isinstance(self.species, CustomSpecies)
            and self.species.can_change_movepool()
            and (movepool := self.species.movepool)
        ):

            doc.add_page_break()
            doc.add_heading("Movepool", level=1)
            if movepool.level:
                doc.add_heading("Level Moves", level=2)
                for k, v in movepool.level.items():
                    if o := ", ".join(x.name for x in sorted(v, key=lambda x: x.name)):
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(f"Level {k:02d}: ").bold = True
                        p.add_run(o)

            if tm := ", ".join(x.name for x in sorted(movepool.tm, key=lambda x: x.name)):
                doc.add_heading("TM Moves", level=2)
                doc.add_paragraph(tm).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if tutor := ", ".join(x.name for x in sorted(movepool.tutor, key=lambda x: x.name)):
                doc.add_heading("Tutor Moves", level=2)
                doc.add_paragraph(tutor).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if egg := ", ".join(x.name for x in sorted(movepool.egg, key=lambda x: x.name)):
                doc.add_heading("Egg Moves", level=2)
                doc.add_paragraph(egg).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if other := ", ".join(x.name for x in sorted(movepool.other, key=lambda x: x.name)):
                doc.add_heading("Other Moves", level=2)
                doc.add_paragraph(other).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(fp := BytesIO())
        fp.seek(0)
        return File(fp=fp, filename=f"{self.id or 'Character'}.docx")

    async def to_pdf(self, bot: CustomBot) -> File:
        """To PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        content: list[Flowable | tuple[Flowable]] = [
            Table(
                [[k, v] for k, v in self.params_header.items()],
                colWidths=[100, 300],
                style=TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.black),
                        ("BOX", (0, 0), (-1, -1), 0.25, colors.black),
                    ]
                ),
            )
        ]

        # Add heading
        content.append(Paragraph(f"<strong>{self.pronoun_emoji}〛{self.name}</strong>", styles["Heading1"]))

        # Add an image
        if img_file := await bot.get_file(self.image_url):
            img_stream = BytesIO(img_file.fp.read())
            content.append(Image(img_stream))

        if self.moveset:
            content.append(Paragraph("Favorite Moves", styles["Heading2"]))
            for item in self.moveset:
                content.append(Paragraph(f"• {item.name} - {item.category.name} - {item.type.name}", styles["Normal"]))

        if self.backstory or self.extra or self.personality:
            content.append(PageBreak())

        if self.backstory:
            content.extend(
                (
                    Spacer(1, 12),
                    Paragraph("Bio", styles["Heading2"]),
                    Paragraph(self.backstory, styles["Normal"]),
                )
            )

        if self.personality:
            content.extend(
                (
                    Spacer(1, 12),
                    Paragraph("Personality", styles["Heading2"]),
                    Paragraph(self.personality, styles["Normal"]),
                )
            )

        if self.extra:
            content.extend(
                (
                    Spacer(1, 12),
                    Paragraph("Extra Information", styles["Heading2"]),
                    Paragraph(self.extra, styles["Normal"]),
                )
            )

        if sp_ability := self.sp_ability:
            content.append(PageBreak())
            content.append(Paragraph(f"Unique Trait: {sp_ability.name}", styles["Heading1"]))
            content.append(Paragraph(sp_ability.kind.phrase, styles["Heading2"]))

            if sp_ability.description:
                content.append(Paragraph(sp_ability.description, styles["Normal"]))

            if sp_ability.pros:
                content.append(Paragraph("How does it make the character's life easier?", styles["Heading3"]))
                content.append(Paragraph(sp_ability.pros, styles["Normal"]))

            if sp_ability.cons:
                content.append(Paragraph("How does it make the character's life harder?", styles["Heading3"]))
                content.append(Paragraph(sp_ability.cons, styles["Normal"]))

            if sp_ability.origin:
                content.append(Paragraph("How was it obtained?", styles["Heading3"]))
                content.append(Paragraph(sp_ability.origin, styles["Normal"]))

        if isinstance(self.species, (Variant, Fakemon)) and (movepool := self.species.movepool):
            content.extend(
                (
                    PageBreak(),
                    Paragraph("Movepool", styles["Heading1"]),
                )
            )

            if movepool.level:
                content.append(Paragraph("Level Moves", styles["Heading2"]))
                content.append(
                    Table(
                        [
                            [f"{k:03d}", ", ".join(x.name for x in sorted(v, key=lambda x: x.name))]
                            for k, v in movepool.level.items()
                            if v
                        ],
                        colWidths=[50, 350],
                        style=TableStyle(
                            [
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.black),
                                ("BOX", (0, 0), (-1, -1), 0.25, colors.black),
                            ]
                        ),
                    )
                )

            if movepool.tm:
                content.extend(
                    (
                        Paragraph("TM Moves", styles["Heading2"]),
                        Paragraph(
                            ", ".join(x.name for x in sorted(movepool.tm, key=lambda x: x.name)),
                            styles["Normal"],
                        ),
                    )
                )

            if movepool.tutor:
                content.extend(
                    (
                        Paragraph("Tutor Moves", styles["Heading2"]),
                        Paragraph(
                            ", ".join(x.name for x in sorted(movepool.tutor, key=lambda x: x.name)),
                            styles["Normal"],
                        ),
                    )
                )

            if movepool.egg:
                content.extend(
                    (
                        Paragraph("Egg Moves", styles["Heading2"]),
                        Paragraph(
                            ", ".join(x.name for x in sorted(movepool.egg, key=lambda x: x.name)),
                            styles["Normal"],
                        ),
                    )
                )

            if movepool.event:
                content.extend(
                    (
                        Paragraph("Event Moves", styles["Heading2"]),
                        Paragraph(
                            ", ".join(x.name for x in sorted(movepool.event, key=lambda x: x.name)),
                            styles["Normal"],
                        ),
                    )
                )

            if movepool.other:
                content.extend(
                    (
                        Paragraph("Other Moves", styles["Heading2"]),
                        Paragraph(
                            ", ".join(x.name for x in sorted(movepool.other, key=lambda x: x.name)),
                            styles["Normal"],
                        ),
                    )
                )

        # Build the PDF document
        doc.build(content)
        buffer.seek(0)

        return File(fp=buffer, filename=f"{self.id or 'Character'}.pdf")

    def generated_image(self, background: Optional[str] = None) -> Optional[str]:
        """Generated Image

        Returns
        -------
        str
            URL
        """
        if isinstance(image := self.image, int):
            return self.image_url
        if image := image or self.default_image:
            if not background:
                background = "background_Y8q8PAtEV.png"
            kit = ImageKit(base=background, width=900, height=450, format="png")
            kit.add_image(image=image, height=450)
            for index, item in enumerate(self.types):
                kit.add_image(image=item.icon, width=200, height=44, x=-10, y=44 * index + 10)
            return kit.url

    @classmethod
    def collage(cls, ocs: Iterable[Character], background: Optional[str] = None, font: bool = True):
        items: list[Character | None] = list(ocs)[:6]

        if background and (amount := 6 - len(items)):
            items.extend([None] * amount)

        kit = ImageKit(base=background or "OC_list_9a1DZPDet.png", width=1500, height=1000, format="png")
        for index, oc in enumerate(items):
            x = 500 * (index % 3) + 25
            y = 500 * (index // 3) + 25
            if oc is None:
                kit.add_image(image="placeholder_uSDglnt-E.png", height=450, width=450, x=x, y=y)
            elif isinstance(oc, Character):
                kit.add_image(image=oc.image_url, height=450, width=450, x=x, y=y)
                for idx, item in enumerate(oc.types):
                    kit.add_image(image=item.icon, width=200, height=44, x=250 + x, y=y + 44 * idx)
                if font:
                    kit.add_text(
                        text=oc.name,
                        width=450,
                        x=x,
                        y=y + 400,
                        background=0xFFFFFF,
                        background_transparency=70,
                        font=Fonts.Whitney_Black,
                        font_size=36,
                        padding=15,
                    )
        return kit.url

    @classmethod
    def rack(cls, ocs: Iterable[Character], font: bool = True):
        items: list[Character | None] = list(ocs)[:4]
        kit = ImageKit(base="Rack_FgmVfIYZs.png", width=2000, height=500, format="png")
        for index, oc in enumerate(items):
            x, y = 500 * index + 25, 25
            if oc is None:
                kit.add_image(image="placeholder_uSDglnt-E.png", height=450, width=450, x=x, y=y)
            elif isinstance(oc, Character):
                kit.add_image(image=oc.image_url, height=450, width=450, x=x, y=y)
                for idx, item in enumerate(oc.types):
                    kit.add_image(image=item.icon, width=200, height=44, x=250 + x, y=y + 44 * idx)
                if font:
                    kit.add_text(
                        text=oc.name,
                        width=450,
                        x=x,
                        y=y + 400,
                        background=0xFFFFFF,
                        background_transparency=70,
                        font=Fonts.Whitney_Black,
                        font_size=36,
                        padding=15,
                    )
        return kit.url

    @classmethod
    def rack2(cls, ocs: Iterable[Character], font: bool = True):
        items: list[Character | None] = list(ocs)[:4]

        kit = ImageKit(base="Rack2_tAEzwkZUI.png", width=1000, height=1000, format="png")
        for index, oc in enumerate(items):
            x = 500 * (index % 2) + 25
            y = 500 * (index // 2) + 25
            if oc is None:
                kit.add_image(image="placeholder_uSDglnt-E.png", height=450, width=450, x=x, y=y)
            elif isinstance(oc, Character):
                kit.add_image(image=oc.image_url, height=450, width=450, x=x, y=y)
                for idx, item in enumerate(oc.types):
                    kit.add_image(image=item.icon, width=200, height=44, x=250 + x, y=y + 44 * idx)
                if font:
                    kit.add_text(
                        text=oc.name,
                        width=450,
                        x=x,
                        y=y + 400,
                        background=0xFFFFFF,
                        background_transparency=70,
                        font=Fonts.Whitney_Black,
                        font_size=36,
                        padding=15,
                    )
        return kit.url

    def __repr__(self) -> str:
        species = self.species.name if self.species else None
        return f"{self.gender.name} - {species}"

    @classmethod
    def process(cls, **kwargs) -> Character:
        """Function used for processing a dict, to a character
        Returns
        -------
        Character
            Character given the paraneters
        """
        data: dict[str, Any] = {k.lower(): v for k, v in kwargs.items()}

        base = common_pop_get(
            data,
            "base",
            "preevo",
            "pre evo",
            "pre_evo",
        )

        if fakemon := data.pop("fakemon", ""):
            name: str = fakemon.title()
            if species := Fakemon.single_deduce(base):
                species.name = name
            else:
                species = Fakemon(name=name)

            data["species"] = species
        elif variant := data.pop("variant", ""):
            if species := Variant.from_base(base=base):
                species.name = variant
            else:
                for item in variant.split(" "):
                    if species := Variant.single_deduce(item):
                        species.name = variant.title()
                        break

            data["species"] = species
        elif (species := Fusion.single_deduce(",".join(data.pop("fusion", [])))) or (
            (aux := common_pop_get(data, "species", "pokemon")) and (species := Species.any_deduce(aux))
        ):
            data["species"] = species

        if isinstance(age := common_pop_get(data, "age", "years"), str):
            data["age"] = AgeGroup.parse(int_check(age))

        if pronoun_info := common_pop_get(data, "pronoun", "gender", "pronouns"):
            data["pronoun"] = Pronoun.deduce_many(pronoun_info)

        data.pop("stats", {})

        if move_info := common_pop_get(data, "moveset", "moves"):
            if isinstance(move_info, str):
                move_info = move_info.split(",")
            move_info = [x for x in move_info if x.lower() != "move"]
            if moveset := Move.deduce_many(*move_info):
                data["moveset"] = moveset

        type_info = common_pop_get(data, "types", "type")
        movepool = Movepool.from_dict(**data.pop("movepool", dict(event=data.get("moveset", set()))))
        data["sp_ability"] = common_pop_get(data, "spability", "sp_ability", "unique_trait")

        if not species:
            data.pop("moveset", None)
        else:
            if type_info and (types := TypingEnum.deduce_many(*type_info)):
                if isinstance(species, (Fusion, CustomSpecies)):
                    species.types = types
                elif species.types != types:
                    types_txt = "/".join(i.name for i in types)
                    species = Variant.from_base(base=species, name=f"{types_txt}-Typed {species.name}")
                    species.types = types

            if isinstance(species, CustomSpecies) and (species.base is None or species.base.movepool != movepool):
                species.movepool = movepool

            data = {k: v for k, v in data.items() if v}
            data["species"] = species

        return cls.from_dict(data)


class CharacterTransform(Transformer):
    async def transform(self, interaction: Interaction[CustomBot], value: str, /):
        db: AsyncIOMotorCollection = interaction.client.mongo_db("Characters")
        if not (member := interaction.namespace.member):
            member = interaction.client.supporting.get(interaction.user, interaction.user)
        ocs = {
            oc.id: oc
            async for x in db.find({"author": member.id, "server": interaction.guild_id})
            if (oc := Character.from_mongo_dict(x))
        }
        if value.isdigit() and (oc := ocs.get(int(value))):
            return oc
        if options := process.extractOne(
            value,
            choices=ocs.values(),
            processor=lambda x: getattr(x, "name", x),
            score_cutoff=60,
        ):
            return options[0]

    async def autocomplete(self, interaction: Interaction[CustomBot], value: str, /) -> list[Choice[str]]:
        db: AsyncIOMotorCollection = interaction.client.mongo_db("Characters")
        if not (member := interaction.namespace.member):
            member = interaction.client.supporting.get(interaction.user, interaction.user)
        ocs = {
            oc.id: oc
            async for x in db.find({"author": member.id, "server": interaction.guild_id})
            if (oc := Character.from_mongo_dict(x))
        }
        if value.isdigit() and (oc := ocs.get(int(value))):
            options = [oc]
        elif options := process.extract(
            value,
            choices=ocs.values(),
            limit=25,
            processor=lambda x: getattr(x, "name", x),
            score_cutoff=60,
        ):
            options = [x[0] for x in options]
        elif not value:
            options = list(ocs.values())[:25]
        return [Choice(name=x.name, value=str(x.id)) for x in options]


CharacterArg = Transform[Character, CharacterTransform]
